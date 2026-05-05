import unittest
import io
from docx import Document
from core.ingestion import DocumentParser


class TestDocumentParser(unittest.TestCase):
    def setUp(self):
        self.parser = DocumentParser()
        doc = Document()
        doc.add_paragraph("This is a test document.")
        doc.add_paragraph("It contains multiple paragraphs.")
        buffer = io.BytesIO()
        doc.save(buffer)
        self.sample_bytes = buffer.getvalue()

    def test_parse_docx_returns_text(self):
        result = self.parser._parse_docx(self.sample_bytes)
        self.assertIsInstance(result, str)
        self.assertIn("This is a test document.", result)


if __name__ == "__main__":
    unittest.main()
